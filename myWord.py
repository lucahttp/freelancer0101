
# document_template = "./template_2.docx"
document_template = "template.docx"
document_template_html = "template.html"



def word_read(word_file):
    import docx
    #doc = docx.Document("E:/my_word_file.docx")
    doc = docx.Document(word_file)
    all_paras = doc.paragraphs
    len(all_paras)
    for para in all_paras:
        print(para.text)
        print("-------")

    import docx
 
    # open connection to Word Document
    doc = docx.Document(word_file)
    
    # read in each paragraph in file
    result = [p.text for p in doc.paragraphs]
    print(result)
    pass


def word_crate(word_file):
    # https://stackabuse.com/reading-and-writing-ms-word-files-in-python-via-python-docx-module/
    # pip install python-docx
    import docx

    mydoc = docx.Document()
    third_para = mydoc.add_paragraph("This is the third paragraph.")
    third_para.add_run(" this is a section at the end of third paragraph")
    #mydoc.save("E:/my_written_file.docx")

    mydoc.add_paragraph("This is first paragraph of a MS Word file.")
    mydoc.add_heading("This is level 1 heading", 0)
    mydoc.add_heading("This is level 2 heading", 1)
    mydoc.add_heading("This is level 3 heading", 2)
    mydoc.save(word_file)
    pass


def word_crate_from_template(word_file):
    # https://docxtpl.readthedocs.io/en/latest/
    # https://blog.formpl.us/how-to-generate-word-documents-from-templates-using-python-cb039ea2c890
    from docxtpl import DocxTemplate

    #doc = DocxTemplate("./inline_image_tpl.docx")
    doc = DocxTemplate("./template.docx")
    context = { 'name' : "luca" ,"greetings":"nice dick bro","unit":"345sdfv","arrival":"14/09/2020","deperture":"15/09/2023","saludation":"hasta la proximaa","contact":"Pam Martin, REALTOR","company_name":"Pam Martin - Keller Williams","phone":"(251)269-8864 or (251)279-0717"}
    doc.render(context)
    #doc.save("generated_doc.docx")
    doc.save(word_file)
    pass

def word_crate_from_template_mail_merge(word_file):
    # https://pbpython.com/python-word-template.html
    # pip install docx-mailmerge
    #from __future__ import print_function
    from mailmerge import MailMerge
    from datetime import date

    template = "template.docx"
    document = MailMerge(template)  
    print(document.get_merge_fields())
    print(document.columns)
    pass





def add_signature(template, context, signature):
    from docx.shared import Cm
    from docxtpl import DocxTemplate, InlineImage
    tmp = DocxTemplate(template)

    img_size = Cm(6.75)  # sets the size of the image
    sign = InlineImage(tmp, signature, img_size)
    context['signature'] = sign  # adds the InlineImage object to the context
    pass



def word_crate_from_template_with_qrcode(word_file,image_qrcode):
    # https://docxtpl.readthedocs.io/en/latest/
    # https://blog.formpl.us/how-to-generate-word-documents-from-templates-using-python-cb039ea2c890
    from docx.shared import Cm
    from docxtpl import DocxTemplate, InlineImage
    #doc = DocxTemplate("./inline_image_tpl.docx")
    doc = DocxTemplate(document_template)
    
    img_size = Cm(6.75)  # sets the size of the image
    sign = InlineImage(doc,image_qrcode, img_size)

    context = { 'name' : "luca" ,"greetings":"nice dick bro","unit":"345sdfv","arrival":"14/09/2020","deperture":"15/09/2023","saludation":"hasta la proximaa","contact":"Pam Martin, REALTOR","company_name":"Pam Martin - Keller Williams","phone":"(251)269-8864 or (251)279-0717"}
    
    context['qrcode'] = sign  # adds the InlineImage object to the context

    doc.render(context)
    #doc.save("generated_doc.docx")
    doc.save(word_file)
    return word_file



""" 
def resource_path(relative_path):
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)

""" 
import os
import sys
from os import chdir
from os.path import join
from os.path import dirname
from os import environ

def resource_path2(relative_path):
    import os
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def resource_path(relative):
    return os.path.join(
        os.environ.get(
            "_MEIPASS2",
            os.path.abspath(".")
        ),
        relative
    )
def word_create_from_template(word_file,image_qrcode,info_name,info_greeting,info_unit,info_arrival,info_deperture,info_saludation,info_contact,info_company_name,info_phone):
    # https://docxtpl.readthedocs.io/en/latest/
    # https://blog.formpl.us/how-to-generate-word-documents-from-templates-using-python-cb039ea2c890
    # from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Cm
    #from html import escape
    from docxtpl import DocxTemplate, InlineImage

    
    filename = 'myfilesname.type'
    filename = document_template
    
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller >= 1.6
        chdir(sys._MEIPASS)
        filename = join(sys._MEIPASS, filename)
    elif '_MEIPASS2' in environ:
        # PyInstaller < 1.6 (tested on 1.5 only)
        chdir(environ['_MEIPASS2'])
        filename = join(environ['_MEIPASS2'], filename)
    else:
        chdir(dirname(sys.argv[0]))
        filename = join(dirname(sys.argv[0]), filename)
        
    # doc = DocxTemplate("./inline_image_tpl.docx")
    # doc = DocxTemplate(resource_path(document_template))
    doc = DocxTemplate(filename)
    import myConfig
    image_qrcode = myConfig.getPath(image_qrcode)
    img_size = Cm(6.75)  # sets the size of the image
    info_qrcode = InlineImage(doc,image_qrcode, img_size)

    #context = { 'name' : "luca" ,"greetings":"nice dick bro","unit":"345sdfv","arrival":"14/09/2020","deperture":"15/09/2023","saludation":"hasta la proximaa","contact":"Pam Martin, REALTOR","company_name":"Pam Martin - Keller Williams","phone":"(251)269-8864 or (251)279-0717"}
    
    context = { 'name' : info_name ,
                "greetings":info_greeting,
                "unit":info_unit,
                "arrival":info_arrival,
                "deperture":info_deperture,
                "saludation":info_saludation,
                "contact":info_contact,
                "company_name":info_company_name,
                "phone":info_phone,
                'qrcode':info_qrcode}
    
    #context['qrcode'] = sign  # adds the InlineImage object to the context

    doc.render(context)
    #doc.save("generated_doc.docx")
    doc.save(word_file)
    return word_file


def create_email_from_html_template(image_qrcode,name,greetings,unit,arrival,deperture,saludation,contact,company_name,phone):

    def getFileContent(file_path):
        """
        with open(file_path, "rb") as file_content:
            return file_content
        """
        f = open(file_path, "r")
        #print(f.read())
        return f.read()

    # html_template_code = getFileContent(document_template_html)
    # getFileContent()
    # txt1 = html_code.format(name = "John", unit = 36)
    """
    attachement = "exe works\qrcode001.png"
    attachement = image_qrcode
    name = "luca"
    greetings = "nice dick bro"
    unit = "345sdfv"
    arrival = "14/09/2020"
    deperture = "15/09/2023"
    saludation = "hasta la proximaa"
    contact = "Pam Martin, REALTOR"
    company_name = "Pam Martin - Keller Williams"
    phone = "(251)269-8864 or (251)279-0717"
    """
    attachement = image_qrcode.split("\\")[-1]
    #attachement = attachement.split(".")[0]
    print(attachement)
    # print()
    
    # qrcode_html = "<img src='cid:"+attachement+"'> "
    # qrcode_html = "<img src="+attachement+">"
    # qrcode_html = '<img src="cid:qrcode001">'
    # image_source = "cid:qrcode001"
    # image_source = attachement
    # image_source = "cid:"+attachement
    # image_source = "cid:qrcode001"
    # qrcode_html = '<img src="'+image_source+'" alt="qrcode" title="QR Code" style="display:block">'
    # https://stackoverflow.com/questions/44544369/i-am-not-able-to-add-an-image-in-email-body-using-python-i-am-able-to-add-a-pi
    import base64
    # https://stackoverflow.com/questions/3715493/encoding-an-image-file-with-base64
    def get_base64_encoded_image(image_path):
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode('utf-8')

    #encoded_image = get_base64_encoded_image(attachement)
    # qrcode_html = '<img src="data:image/png;base64,%s"/>' % get_base64_encoded_image(attachement)
    # qrcode_html = '<img src="data:image/jpeg;base64,%s"  alt="qrcode" title="QR Code" style="display:block;"  width="350" height="350" style="text-align:center;"/>' % get_base64_encoded_image(attachement)
    # qrcode_html = '<img src="data:image/jpeg;base64,%s"  alt="qrcode" title="QR Code" width="350" height="350" />' % get_base64_encoded_image(attachement)

    # qrcode_html = '<img src="'+attachement+'" alt="imagename" class="imageclass" />'
    # https://stackoverflow.com/questions/21437916/html-image-not-showing-in-gmail#31820008
    # qrcode_html = '<div style="height: 24px; width: 24px; display: block; background: url('+attachement+'); background-size: contain;"></div>'
    qrcode_html = '<img src="cid:MyId1">'

    # encoded_image = base64.b64encode(attachement.getvalue()).decode("utf-8")
    # qrcode_html = '<img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAACAAAAAgCAYAAABzenr0AAACR0lEQVRYha1XvU4bQRD+bF/JjzEnpUDwCPROywPgB4h0PUWkFEkLposUIYyEU4N5AEpewnkDCiQcjBQpWLiLjk3DrnZnZ3buTv4ae25mZ+Z2Zr7daxljDGpg++Mv978Y5Nhc6+Di5tk9u7/bR3cjY9eOJnMUh3mg5y0roBjk+PF1F+1WCwCCJKTgpz9/ozjMg+ftVQQ/PtrB508f1OAcau8ADW5xfLRTOzgAZMPxTNy+YpDj6vaPGtxPgvpL7QwAtKXts8GqBveT8P1p5YF5x8nlo+n1p6bXn5ov3x9M+fZmjDGRXBXWH5X/Lv4FdqCLaLAmwX1/VKYJtIwJeYDO+dm3PSePJnO8vJbJhqN62hOUJ8QpoD1Au5kmIentr9TobAK04RyJEOazzjV9KokogVRwjvm6652kniYRJUBrTkft5bUEAGyuddzz7noHALBYls5O09skaE+4HdAYruobUz1FVI6qcy7xRFW95A915pzjiTp6zj7za6fB1lay1/Ssfa8/jRiLw/n1k9tizl7TS/aZ3xDakdqUByR/gDcF0qJV8QAXHACy+7v9wGA4ngWLVskDo8kcg4Ot8FpGa8PV0I7MyeWjq53f7Zrer3nyOLYJpJJowgN+g9IExNNQ4vLFskwyJtVrd8JoB7g3b4rz66dIpv7UHqg611xw/0om8QT7XXBx84zheCbKGui2U9n3p/YAlSVyqRqc+kt+mCyWJTSeoMGjOQciOQDXA6kjVTsL6JhpYHtA+wihPaGOWgLqnVACPQua4j8NK7bPLP4+qQAAAABJRU5ErkJggg==" width="32" height="32">'
    """
    html_code = html_template_code.format(name  =  name ,
                    greetings = greetings,
                    unit = unit,
                    arrival = arrival,
                    deperture = deperture,
                    saludation = saludation,
                    contact = contact,
                    company_name = company_name,
                    phone = phone,
                    qrcode = qrcode_html)
    """
    from string import Template
    html_template_code = getFileContent(document_template_html)

    html_template_code = html_template_code.replace("{","$")
    html_template_code = html_template_code.replace("}","")
    #s = Template('$who likes $what')
    #s = s.substitute(who='tim', what='kung pao')
    s = Template(html_template_code)
    html_completed = s.substitute(name = name ,
                greetings = greetings,
                unit = unit,
                arrival = arrival,
                deperture = deperture,
                saludation = saludation,
                contact = contact,
                company_name = company_name,
                phone = phone,
                qrcode = qrcode_html)
    # print(txt1)
    Html_file= open("index.html","w")
    Html_file.write(html_completed)
    Html_file.close()
    print(html_completed)
    return html_completed


template_code ="""<p style='margin-top:0cm;margin-right:0cm;margin-bottom:8.0pt;margin-left:0cm;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:24px;line-height:107%;font-family:"Perpetua",serif;'>{name}</span><span style='font-size:24px;line-height:107%;font-family:"Perpetua",serif;'>, {greetings}</span></p>
<p style='margin-top:0cm;margin-right:0cm;margin-bottom:8.0pt;margin-left:0cm;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:24px;line-height:107%;font-family:"Perpetua",serif;'>Relax! You&rsquo;ll be on the beach soon!</span></p>
<p style='margin-top:0cm;margin-right:0cm;margin-bottom:8.0pt;margin-left:0cm;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><strong><span style='font-size:32px;line-height:107%;font-family:"Perpetua",serif;'>Welcome to The Beach Club!</span></strong></p>
<p style='margin-top:0cm;margin-right:0cm;margin-bottom:8.0pt;margin-left:0cm;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:21px;line-height:107%;font-family:"Perpetua",serif;'>We&rsquo;ll have everything ready for you at 4 PM on {deperture}</span><span style='font-size:24px;line-height:107%;font-family:"Perpetua",serif;'>.</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><strong><span style='font-size:19px;font-family:"Perpetua",serif;'>Please have this email and your Driver&rsquo;s License out and ready to present to the attendant at the gate. They will scan the code below and your driver&rsquo;s license to confirm your reservation of:</span></strong></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><strong><span style='font-size:11px;font-family:"Perpetua",serif;'>&nbsp;</span></strong></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><strong><span style='font-size:21px;font-family:"Perpetua",serif;'>Unit {unit}&nbsp;from {arrival}&nbsp;&ndash; {deperture}.</span></strong></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><strong><span style='font-size:11px;font-family:"Perpetua",serif;'>&nbsp;</span></strong></p>
<p style='margin-top:0cm;margin-right:-9.0pt;margin-bottom:8.0pt;margin-left:-22.5pt;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:19px;line-height:107%;font-family:"Perpetua",serif;'>Once you have registered at the gate you may use either of the entrance gates for the remainder of your stay. The system will recognize your vehicle automatically and open the gate.</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:11px;font-family:"Perpetua",serif;'>{qrcode}</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:11px;font-family:"Perpetua",serif;'>&nbsp;</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:11px;font-family:"Perpetua",serif;'>&nbsp;</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:19px;font-family:"Perpetua",serif;'>Thank you for your help to expedite the check in process.</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:19px;font-family:"Perpetua",serif;'>Soon you will be on the beach enjoying our beautiful resort.</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:11px;font-family:"Perpetua",serif;'>&nbsp;</span></p>
<p style='margin-top:0cm;margin-right:0cm;margin-bottom:8.0pt;margin-left:0cm;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;text-align:center;'><span style='font-size:19px;line-height:107%;font-family:"Perpetua",serif;'>{saludation}</span></p>
<p style='margin-top:0cm;margin-right:0cm;margin-bottom:8.0pt;margin-left:0cm;line-height:107%;font-size:15px;font-family:"Calibri",sans-serif;'><span style='font-size:19px;line-height:107%;font-family:"Perpetua",serif;'>&nbsp;</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;'><span style='font-size:19px;font-family:"Perpetua",serif;'>{contact}</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;'><span style='font-size:19px;font-family:"Perpetua",serif;'>{company_name}&nbsp;</span></p>
<p style='margin:0cm;font-size:15px;font-family:"Calibri",sans-serif;'><span style='font-size:19px;font-family:"Perpetua",serif;'>{phone}</span></p>"""